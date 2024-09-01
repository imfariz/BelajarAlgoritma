from docx import Document
from docx.shared import Inches

document = Document("./TemplateTitle.docx")

def templating(data):
    table = document.add_table(rows=4, cols=1, style="CustomTable")
    table.rows[0].cells[0].text = data['spm']
    table.rows[0].height = Inches(0.6)
    table.rows[1].cells[0].text = data['cv']
    table.rows[1].height = Inches(0.4)
    table.rows[2].cells[0].text = data['paket']
    table.rows[2].height = Inches(0.8)
    table.rows[3].cells[0].text = data['nilai']
    table.rows[3].height = Inches(0.6)
    document.add_paragraph('')

def converting(data):
    for x in range(len(data)):
        templating(data[x])
        if(x%3 == 0):
            document.add_page_break()
    document.save(f"test.docx")
